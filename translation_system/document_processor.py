"""
نظام الترجمة المكتبي - معالج الوثائق
Translation Office System - Document Processor
"""

import os
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import PyPDF2
import io
import re

from config import (
    PDF_MARGIN_TOP,
    PDF_MARGIN_BOTTOM,
    PDF_MARGIN_LEFT,
    PDF_MARGIN_RIGHT,
    SUPPORTED_FILE_TYPES
)


class DocumentProcessor:
    """معالج الوثائق للتعامل مع ملفات Word وPDF"""
    
    def __init__(self):
        self.supported_extensions = ['.docx', '.pdf']
    
    def read_document(self, file_path: str) -> Tuple[bool, str, str]:
        """
        قراءة محتوى الوثيقة
        Returns: (success, content, file_type)
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return False, "", ""
            
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.docx':
                return self._read_docx(file_path)
            elif file_extension == '.pdf':
                return self._read_pdf(file_path)
            else:
                return False, "", f"نوع الملف غير مدعوم: {file_extension}"
                
        except Exception as e:
            return False, "", f"خطأ في قراءة الملف: {str(e)}"
    
    def _read_docx(self, file_path: Path) -> Tuple[bool, str, str]:
        """قراءة ملف Word"""
        try:
            doc = Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            return True, '\n'.join(content), 'docx'
            
        except Exception as e:
            return False, "", f"خطأ في قراءة ملف Word: {str(e)}"
    
    def _read_pdf(self, file_path: Path) -> Tuple[bool, str, str]:
        """قراءة ملف PDF"""
        try:
            content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text and text.strip():
                        content.append(text)
            
            return True, '\n'.join(content), 'pdf'
            
        except Exception as e:
            return False, "", f"خطأ في قراءة ملف PDF: {str(e)}"
    
    def create_translation_template(self, original_content: str, output_path: str) -> bool:
        """إنشاء نموذج ترجمة من المحتوى الأصلي"""
        try:
            doc = Document()
            
            # إعداد الهوامش
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(PDF_MARGIN_TOP)
                section.bottom_margin = Cm(PDF_MARGIN_BOTTOM)
                section.left_margin = Cm(PDF_MARGIN_LEFT)
                section.right_margin = Cm(PDF_MARGIN_RIGHT)
            
            # إضافة العنوان
            title = doc.add_heading('نموذج الترجمة', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # إضافة المحتوى الأصلي
            doc.add_heading('المحتوى الأصلي:', level=1)
            
            # تقسيم المحتوى إلى فقرات
            paragraphs = original_content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(para_text)
            
            # إضافة مساحة للترجمة
            doc.add_page_break()
            doc.add_heading('الترجمة:', level=1)
            
            # إضافة فقرات فارغة للترجمة
            for para_text in paragraphs:
                if para_text.strip():
                    paragraph = doc.add_paragraph()
                    paragraph.add_run('_' * 50)  # خطوط للترجمة
            
            # حفظ الملف
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"خطأ في إنشاء نموذج الترجمة: {e}")
            return False
    
    def create_final_document(self, project_data: Dict[str, Any], output_path: str) -> bool:
        """إنشاء الوثيقة النهائية مع الترجمة والمستند الأصلي"""
        try:
            doc = Document()
            
            # إعداد الهوامش
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(PDF_MARGIN_TOP)
                section.bottom_margin = Cm(PDF_MARGIN_BOTTOM)
                section.left_margin = Cm(PDF_MARGIN_LEFT)
                section.right_margin = Cm(PDF_MARGIN_RIGHT)
            
            # إضافة شعار المكتب (إذا كان موجوداً)
            # self._add_company_logo(doc)
            
            # إضافة العنوان
            title = doc.add_heading('وثيقة الترجمة الرسمية', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # إضافة معلومات المشروع
            self._add_project_info(doc, project_data)
            
            # إضافة الترجمة
            doc.add_page_break()
            doc.add_heading('الترجمة:', level=1)
            
            translated_content = project_data.get('translated_content', '')
            if translated_content:
                paragraphs = translated_content.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        paragraph = doc.add_paragraph()
                        paragraph.add_run(para_text)
            
            # إضافة المستند الأصلي
            doc.add_page_break()
            doc.add_heading('المستند الأصلي:', level=1)
            
            original_content = project_data.get('original_content', '')
            if original_content:
                paragraphs = original_content.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        paragraph = doc.add_paragraph()
                        paragraph.add_run(para_text)
            
            # إضافة صندوق اعتماد المترجم
            doc.add_page_break()
            self._add_translator_certification_box(doc, project_data)
            
            # حفظ الملف
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"خطأ في إنشاء الوثيقة النهائية: {e}")
            return False
    
    def _add_project_info(self, doc: Document, project_data: Dict[str, Any]):
        """إضافة معلومات المشروع"""
        # معلومات العميل
        client_info = doc.add_paragraph()
        client_info.add_run(f"اسم العميل: {project_data.get('client_name', '')}").bold = True
        client_info.add_run(f"\nالبريد الإلكتروني: {project_data.get('client_email', '')}")
        
        # معلومات الترجمة
        translation_info = doc.add_paragraph()
        translation_info.add_run(f"اللغة المصدر: {project_data.get('source_language', '')}").bold = True
        translation_info.add_run(f"\nاللغة الهدف: {project_data.get('target_language', '')}")
        translation_info.add_run(f"\nتاريخ الترجمة: {project_data.get('translation_date', '')}")
        
        # معلومات المترجم
        translator_info = doc.add_paragraph()
        translator_info.add_run(f"المترجم: {project_data.get('translator_name', '')}").bold = True
        translator_info.add_run(f"\nرقم الترخيص: {project_data.get('translator_license', '')}")
    
    def _add_translator_certification_box(self, doc: Document, project_data: Dict[str, Any]):
        """إضافة صندوق اعتماد المترجم"""
        # عنوان الصندوق
        title = doc.add_heading('اعتماد المترجم المحلف', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # إطار الصندوق
        certification_box = doc.add_paragraph()
        certification_box.add_run("أعتمد أنا الموقع أدناه، المترجم المحلف، أن الترجمة المرفقة صحيحة ومطابقة للمستند الأصلي.")
        
        # معلومات المترجم
        translator_details = doc.add_paragraph()
        translator_details.add_run(f"اسم المترجم: {project_data.get('translator_name', '')}").bold = True
        translator_details.add_run(f"\nرقم الترخيص: {project_data.get('translator_license', '')}")
        translator_details.add_run(f"\nاللغة المصدر: {project_data.get('source_language', '')}")
        translator_details.add_run(f"\nاللغة الهدف: {project_data.get('target_language', '')}")
        translator_details.add_run(f"\nتاريخ الاعتماد: {project_data.get('certification_date', '')}")
        
        # مساحة للتوقيع
        signature_space = doc.add_paragraph()
        signature_space.add_run("\n\n\n")
        signature_space.add_run("التوقيع: _________________").bold = True
        
        # إضافة QR Code (سيتم إضافته لاحقاً)
        qr_note = doc.add_paragraph()
        qr_note.add_run("\nQR Code للتحقق من صحة الوثيقة")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """استخراج النص من ملف Word"""
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            return '\n'.join(text_content)
            
        except Exception as e:
            print(f"خطأ في استخراج النص من ملف Word: {e}")
            return ""
    
    def merge_documents(self, doc1_path: str, doc2_path: str, output_path: str) -> bool:
        """دمج وثيقتين Word"""
        try:
            doc1 = Document(doc1_path)
            doc2 = Document(doc2_path)
            
            # إضافة فاصل الصفحة
            doc1.add_page_break()
            
            # نسخ جميع الفقرات من الوثيقة الثانية
            for element in doc2.element.body:
                doc1.element.body.append(element)
            
            # حفظ الوثيقة المدمجة
            doc1.save(output_path)
            return True
            
        except Exception as e:
            print(f"خطأ في دمج الوثائق: {e}")
            return False
    
    def validate_file_type(self, file_path: str) -> bool:
        """التحقق من نوع الملف"""
        file_extension = Path(file_path).suffix.lower()
        return file_extension in self.supported_extensions
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """الحصول على معلومات الملف"""
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return {}
            
            file_info = {
                'name': file_path.name,
                'size': file_path.stat().st_size,
                'extension': file_path.suffix.lower(),
                'created': file_path.stat().st_ctime,
                'modified': file_path.stat().st_mtime
            }
            
            # إضافة معلومات خاصة بنوع الملف
            if file_info['extension'] == '.docx':
                doc = Document(file_path)
                file_info['paragraphs'] = len(doc.paragraphs)
                file_info['pages'] = len(doc.sections)
            elif file_info['extension'] == '.pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    file_info['pages'] = len(pdf_reader.pages)
            
            return file_info
            
        except Exception as e:
            print(f"خطأ في الحصول على معلومات الملف: {e}")
            return {}
